from docx import Document

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = ''
    # 提取段落文本
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    # 提取表格文本，保持表格形式
    for table in doc.tables:
        for row in table.rows:
            row_text = ''
            for cell in row.cells:
                row_text += cell.text + ' '  # 使用空格而不是换行连接单元格文本
            text += row_text.strip() + '\n'  # 移除行尾多余的空格，并添加换行符
    return text

docx_path = r'C:\Users\lyz13\OneDrive\CloudyLake Programming\MeteoStation of CloudyLake\MeteoStation-of-CloudyLake\station info.docx'
extracted_text = extract_text_from_docx(docx_path)
print(extracted_text)
txt_path = r'C:\Users\lyz13\OneDrive\CloudyLake Programming\MeteoStation of CloudyLake\MeteoStation-of-CloudyLake\extracted_text.txt'
with open(txt_path, 'w', encoding='utf-8') as file:
    file.write(extracted_text)
print("Text extracted from docx file has been saved as a txt file.")